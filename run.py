# Load python packages
from alive_progress import alive_it, config_handler
from bs4 import BeautifulSoup
from docx import Document
from pathlib import Path
from tabulate import tabulate
import json
import multiprocessing as mp
import pandas as pd
import requests

# Define user agent
USER_AGENT = 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/103.0.0.0 Safari/537.36'

# Configure loading bar
config_handler.set_global(
    length=75,
    spinner='classic',
    bar='classic2',
    force_tty=True,
    dual_line=True
)

# Create json file
def create_json():
    # Set up image path
    image_path = Path.cwd() / 'image'
    if not image_path.is_dir():
        image_path.mkdir()

    # Read word document
    file_path = Path.cwd() / 'workout.docx'
    document = Document(file_path)

    # Extract titles
    paragraphs = document.paragraphs
    titles = [x.text for x in paragraphs if x.text.strip()]

    # Extract data from tables
    tables = document.tables
    keys = ['day1', 'day2', 'day3', 'day4', 'day5', 'day6']
    data = {}
    for index, table in enumerate(tables):
        print(titles[index])
        df = create_df(table)
        df['Image'] = get_images(df)
        key = keys[index]
        data[key] = df.to_dict('records')

    # Export json file
    json_path = Path.cwd() / 'data.json'
    with open(json_path, 'w') as file:
        json.dump(data, file)

# Create dataframe from table
def create_df(table):
    num_row = len(table.rows)
    num_col = len(table.rows[0].cells)
    for i in range(num_row):
        row = table.rows[i]
        if i == 0:
            columns = [x.text.strip() for x in row.cells]
            data = []
        else:
            items = []
            for j in range(num_col):
                item = row.cells[j].text.strip()
                items.append(item)
            data.append(items)
    df = pd.DataFrame(data, columns=columns)
    return df

# Download images for all exercises
def get_images(df):
    links = list(df['Link'])
    processes = mp.cpu_count() - 1
    with mp.Pool(processes) as pool:
        results = alive_it(
            pool.imap(get_image, links),
            len(links)
        )
        images = []
        for index, result in enumerate(results):
            results.text(f'Download: {links[index]}')
            images.append(result)
        return images

# Download image
def get_image(url):
    # Create filename from url
    filename = url.split('/')[-2] + '.svg'

    # Get image data
    url = get_link(url)
    headers = {
        'user-agent': USER_AGENT
    }
    response = requests.get(url, headers=headers)

    # Export image file
    image = response.content
    image_path = Path('image') / filename
    with open(image_path, 'wb') as file:
        file.write(image)
    return image_path.as_posix()

# Get image link from website
def get_link(url):
    headers = {
        'user-agent': USER_AGENT
    }
    response = requests.get(url, headers=headers)
    soup = BeautifulSoup(response.text, 'html.parser')
    link = soup.find('img', {'class': 'exImg'})['data-url_male']
    return link

# Create html file
def create_html():
    # Read json data
    json_path = Path.cwd() / 'data.json'
    with open(json_path, 'r') as file:
        data = json.load(file)

    # Read template html
    html_path = Path.cwd() / 'template.html'
    with open(html_path, 'r') as file:
        soup = BeautifulSoup(file, 'html.parser')

    # Create tables for all tabs
    div_list = soup.find_all('div', {'class': 'overflow-auto'})
    for div in div_list:
        div_id = div.parent['id']
        table = create_table(data[div_id])
        div.append(table)

    # Add attributes to table elements
    for table in soup.find_all('table'):
        table['class'] = 'table table-bordered table-hover'

    for thead in soup.find_all('thead'):
        thead['class'] = 'table-dark'

    for th in soup.find_all('th'):
        th['scope'] = 'col'
        th['class'] = 'align-middle'

    # Export to html file
    output_path = Path.cwd() / 'index.html'
    with open(output_path, 'w') as file:
        html = soup.prettify()
        file.write(html)

# Create <table> tag element
def create_table(array):
    table = []
    for item in array:
        row = [
            item['No'],
            create_link(item['Exercise'], item['Link']),
            create_img(item['Image']),
            item['Part']
        ]
        table.append(row)
    headers = ['No', 'Exercise', 'Diagram', 'Part']
    html = tabulate(table, headers, tablefmt='unsafehtml', disable_numparse=True)
    soup = BeautifulSoup(html, 'html.parser')
    return soup

# Create <a> tag element
def create_link(name, url):
    soup = BeautifulSoup(features='html.parser')
    a = soup.new_tag('a')
    a.string = name
    a['href'] = url
    return str(a)

# Create <img> tag element
def create_img(url):
    soup = BeautifulSoup(features='html.parser')
    img = soup.new_tag('img')
    img['src'] = url
    img['alt'] = url.split('/')[-1]
    return str(img)

# Run main program
if __name__ == '__main__':
    create_json()
    create_html()